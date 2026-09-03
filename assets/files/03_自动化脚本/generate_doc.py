#!/usr/bin/env python3
"""
将"非常深科技半年度拍摄规划" Markdown 转为中国党政机关公文格式 Word 文档。
标准：GB/T 9704-2012 党政机关公文格式
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# 配置
# ============================================================
MD_PATH = "/Users/huangshifu/Desktop/claude code/output/非常深科技半年度拍摄规划-2026年6月.md"
OUT_PATH = "/Users/huangshifu/Desktop/claude code/output/非常深科技半年度拍摄规划-2026年6月.docx"

# 公文标准字体（优先使用，不存在则回退）
TITLE_FONT = "方正小标宋简体"          # 标题
TITLE_FONT_FALLBACK = "宋体"
H1_FONT = "黑体"                       # 一级标题
H2_FONT = "楷体_GB2312"               # 二级标题
H2_FONT_FALLBACK = "楷体"
BODY_FONT = "仿宋_GB2312"             # 正文
BODY_FONT_FALLBACK = "仿宋"

# 字号
TITLE_SIZE = Pt(22)                    # 二号
HEADING_SIZE = Pt(16)                  # 三号
BODY_SIZE = Pt(16)                     # 三号
PAGE_NUM_SIZE = Pt(14)                 # 四号

# 页边距 (GB/T 9704)
MARGIN_TOP = Cm(3.7)
MARGIN_BOTTOM = Cm(3.5)
MARGIN_LEFT = Cm(2.8)
MARGIN_RIGHT = Cm(2.6)

# 行距
LINE_SPACING = Pt(28)

# ============================================================
# 辅助函数
# ============================================================
def set_font(run, font_name, font_name_fallback, size, bold=False):
    """设置字体、字号、加粗"""
    run.font.size = size
    run.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    # 设置西文字体回退
    if font_name_fallback:
        run.font.name = font_name_fallback

def set_line_spacing(paragraph, spacing=LINE_SPACING):
    """设置固定行距 28pt"""
    pf = paragraph.paragraph_format
    pf.line_spacing = spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    from docx.enum.text import WD_LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY

def add_title(doc, text):
    """添加公文标题：方正小标宋 22pt 居中"""
    p = doc.add_paragraph()
    set_line_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_font(run, TITLE_FONT, TITLE_FONT_FALLBACK, TITLE_SIZE, bold=False)
    return p

def add_h1(doc, text):
    """一级标题：黑体 16pt 加粗"""
    p = doc.add_paragraph()
    set_line_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_font(run, H1_FONT, None, HEADING_SIZE, bold=True)
    return p

def add_h2(doc, text):
    """二级标题：楷体 16pt 加粗"""
    p = doc.add_paragraph()
    set_line_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_font(run, H2_FONT, H2_FONT_FALLBACK, HEADING_SIZE, bold=True)
    return p

def add_body(doc, text):
    """正文：仿宋 16pt，首行缩进2字符"""
    p = doc.add_paragraph()
    set_line_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # 首行缩进2字符 ≈ 32pt (16pt × 2)
    p.paragraph_format.first_line_indent = Pt(32)
    run = p.add_run(text)
    set_font(run, BODY_FONT, BODY_FONT_FALLBACK, BODY_SIZE, bold=False)
    return p

def add_body_no_indent(doc, text, bold=False, font_name=None, fallback=None):
    """正文无缩进（用于作者署名、日期等）"""
    p = doc.add_paragraph()
    set_line_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    fn = font_name or BODY_FONT
    fb = fallback or BODY_FONT_FALLBACK
    set_font(run, fn, fb, BODY_SIZE, bold=bold)
    return p

def add_blank_line(doc):
    """添加空行"""
    p = doc.add_paragraph()
    set_line_spacing(p)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run("")
    set_font(run, BODY_FONT, BODY_FONT_FALLBACK, BODY_SIZE)
    return p

# ============================================================
# 表格处理
# ============================================================
def add_table_from_md(doc, markdown_table_lines):
    """将 Markdown 表格行转为 Word 表格"""
    # 解析
    rows = []
    for line in markdown_table_lines:
        line = line.strip()
        if not line or line.startswith("|--") or line.startswith("|-"):
            continue
        if line.startswith("|"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            rows.append(cells)

    if len(rows) < 2:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            # 清除默认段落
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            set_line_spacing(p)
            run = p.add_run(cell_text)
            if i == 0:
                # 表头
                set_font(run, H1_FONT, None, Pt(12), bold=True)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                set_font(run, BODY_FONT, BODY_FONT_FALLBACK, Pt(12), bold=False)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)

    # 表后空行
    add_blank_line(doc)


def add_page_numbers(doc):
    """添加页码：页面底部居中，四号半角宋体"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)

        # 添加页码域
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)

        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run._r.append(instrText)

        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run._r.append(fldChar2)

        set_font(run, "宋体", None, PAGE_NUM_SIZE, bold=False)


def set_page_margins(doc):
    """设置 A4 纸张及页边距"""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT


# ============================================================
# 主解析逻辑
# ============================================================
def parse_and_build(md_path, doc):
    """逐行解析 Markdown 并构建 Word 文档"""

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    table_buffer = []
    in_table = False
    in_code_block = False
    separator_count = 0  # 追踪 "---" 水平线

    while i < len(lines):
        line = lines[i].rstrip()

        # 跳过代码块
        if line.startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            i += 1
            continue

        # 处理水平线 "---"
        if re.match(r'^---+$', line.strip()):
            separator_count += 1
            if separator_count == 1:
                # 第一条水平线：分隔标题与正文，跳过
                pass
            elif separator_count == 2:
                # 文档末尾的分隔线，跳过
                pass
            i += 1
            continue

        # 空行
        if not line.strip():
            if in_table:
                add_table_from_md(doc, table_buffer)
                table_buffer = []
                in_table = False
            i += 1
            continue

        # 表格行
        if line.strip().startswith("|") and ("|" in line.strip()[1:]):
            if not in_table:
                in_table = True
                table_buffer = []
            table_buffer.append(line)
            i += 1
            continue

        # 如果之前在处理表格，现在表格结束
        if in_table:
            add_table_from_md(doc, table_buffer)
            table_buffer = []
            in_table = False

        # ---- 标题 # ----
        m = re.match(r'^#\s+(.+)', line)
        if m:
            title_text = m.group(1).strip()
            # 跳过与文件名相同的标题（公文标题已在最前面）
            if "非常深科技" in title_text and "规划" in title_text:
                i += 1
                continue
            add_h1(doc, title_text)
            i += 1
            continue

        # ---- 二级标题 ## ----
        m = re.match(r'^##\s+(.+)', line)
        if m:
            h2_text = m.group(1).strip()
            add_h1(doc, h2_text)  # 公文用一级标题格式
            i += 1
            continue

        # ---- 三级标题 ### ----
        m = re.match(r'^###\s+(.+)', line)
        if m:
            h3_text = m.group(1).strip()
            add_h2(doc, h3_text)
            i += 1
            continue

        # ---- 四级标题 #### ----
        m = re.match(r'^####\s+(.+)', line)
        if m:
            h4_text = m.group(1).strip()
            add_h2(doc, h4_text)
            i += 1
            continue

        # ---- 带数字编号的标题（一）、（二）等 ----
        m = re.match(r'^（([一二三四五六七八九十百]+)）(.+)', line)
        if m:
            add_h2(doc, line.strip())
            i += 1
            continue

        # ---- "一、" 类一级标题 ----
        m = re.match(r'^(一、|二、|三、|四、|五、|六、|七、|八、|九、|十、)(.+)', line)
        if m:
            add_h1(doc, line.strip())
            i += 1
            continue

        # ---- "1. " 数字列表 ----
        m = re.match(r'^(\d+)\.\s+(.+)', line)
        if m:
            # 数字列表作为正文处理（加粗数字部分）
            add_body(doc, line.strip())
            i += 1
            continue

        # ---- "- " 项目符号 ----
        m = re.match(r'^-\s+(.+)', line)
        if m:
            add_body(doc, "— " + m.group(1).strip())
            i += 1
            continue

        # ---- "> " 引用块 ----
        m = re.match(r'^>\s+(.+)', line)
        if m:
            add_body(doc, m.group(1).strip())
            i += 1
            continue

        # ---- "**...**" 加粗段落 ----
        if line.startswith("**") and line.endswith("**"):
            content = line.strip("*")
            add_body_no_indent(doc, content, bold=True)
            i += 1
            continue

        # ---- 普通正文 ----
        # 移除 Markdown 行内格式
        clean = line.strip()
        # 移除加粗标记 **text**
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', clean)
        # 移除行内代码 `text`
        clean = re.sub(r'`([^`]+)`', r'\1', clean)
        # 移除链接 [text](url)
        clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', clean)

        if clean:
            add_body(doc, clean)

        i += 1

    # 处理最后的表格
    if in_table:
        add_table_from_md(doc, table_buffer)


# ============================================================
# 主入口
# ============================================================
def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    style._element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # 设置页边距
    set_page_margins(doc)

    # ==================== 红头文件头部 ====================
    # 发文机关标志（红头大字，居中）
    p = doc.add_paragraph()
    set_line_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run("DeepTech 深科技内容中心")
    set_font(run, TITLE_FONT, "宋体", Pt(18), bold=True)
    run.font.color.rgb = parse_xml(f'<w:color {nsdecls("w")} w:val="FF0000"/>').find(
        qn('w:color')).get(qn('w:val')) if False else None
    # 用红色
    from docx.shared import RGBColor
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    # 红色分隔线
    p = doc.add_paragraph()
    set_line_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    # 添加红色下划线模拟红头分隔线
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="FF0000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run("")
    set_font(run, BODY_FONT, BODY_FONT_FALLBACK, Pt(6))

    add_blank_line(doc)

    # ==================== 密级和编号 ====================
    p = doc.add_paragraph()
    set_line_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run("密级：内部参考")
    set_font(run, BODY_FONT, BODY_FONT_FALLBACK, BODY_SIZE)
    p = doc.add_paragraph()
    set_line_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run("编号：DT-RPT-2026-003")
    set_font(run, BODY_FONT, BODY_FONT_FALLBACK, BODY_SIZE)

    add_blank_line(doc)

    # ==================== 公文标题 ====================
    # 主标题
    add_title(doc, "关于“非常深科技”视频号")
    add_title(doc, "半年度拍摄规划及发展战略的报告")

    add_blank_line(doc)

    # ==================== 主送机关 ====================
    add_body_no_indent(doc, "公司董事会、管理层：", bold=False)

    add_blank_line(doc)

    # ==================== 正文：引语 ====================
    intro = (
        "“非常深科技”系 DeepTech 深科技旗下聚焦硬核科研实景访谈的垂直科技视频 IP，"
        "自二〇二六年四月上线以来，已完成二十余个科研团队的拍摄工作，全网累计播放量突破一百万次，"
        "覆盖微信视频号、抖音、小红书三大主流平台。为系统推进该 IP 的内容建设与商业化进程，"
        "特编制半年度拍摄规划及发展战略，现报告如下。"
    )
    add_body(doc, intro)

    add_blank_line(doc)

    # ==================== 解析 Markdown 正文 ====================
    parse_and_build(MD_PATH, doc)

    # ==================== 文尾 ====================
    add_blank_line(doc)
    add_blank_line(doc)

    # 编制单位
    add_body_no_indent(doc, "编制单位：DeepTech 深科技内容中心", bold=False)
    add_body_no_indent(doc, "联 系 人：黄师傅", bold=False)

    add_blank_line(doc)

    # 日期
    p = doc.add_paragraph()
    set_line_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run("二〇二六年六月九日")
    set_font(run, BODY_FONT, BODY_FONT_FALLBACK, BODY_SIZE)

    add_blank_line(doc)

    # 印发说明
    add_body_no_indent(doc, "（此报告仅供内部决策参考，未经授权不得对外传播）", bold=False)

    # ==================== 页码 ====================
    add_page_numbers(doc)

    # ==================== 保存 ====================
    doc.save(OUT_PATH)
    print(f"✅ 公文已生成：{OUT_PATH}")


if __name__ == "__main__":
    main()
